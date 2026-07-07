# -*- coding: utf-8 -*-
"""report.py — 투자심리검사 개인 진단 리포트(docx) 생성. 명의 '투자농장'."""
from __future__ import annotations
from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xC0, 0x39, 0x2B)


def _kr(run, name="맑은 고딕"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _p(doc, text, size=10.5, bold=False, color=None, align=None, after=3):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    _kr(r)
    return p


def make_docx(prof: dict, extras: dict, radar_png: bytes | None = None) -> bytes:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    _p(doc, "투자 심리검사 — 개인 행동진단 리포트", size=18, bold=True,
       color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _p(doc, f"진단일 {date.today()} · 정리 투자농장", size=10, color=GRAY,
       align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    _p(doc, f"■ 당신의 유형:  {prof['type']}", size=14, bold=True, color=RED)
    _p(doc, f"매매 {prof['n_sells']}회 · 평균 보유 {prof['avg_hold']:.1f}주", size=10,
       color=GRAY)

    # 5가지 매매 습관
    _p(doc, "■ 나의 5가지 매매 습관 점수 (0=없음 … 100=강함)", size=13, bold=True,
       color=NAVY, after=4)
    for k, v in prof["traits"].items():
        bar = "█" * round(v / 10) + "░" * (10 - round(v / 10))
        _p(doc, f"  {k:<9} {bar}  {v:.0f}", size=10.5, after=1)

    if radar_png:
        doc.add_picture(BytesIO(radar_png))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 상황별 '오른 건 팔고 내린 건 버티기' 정도
    _p(doc, "■ 상황별 '오른 건 팔고 내린 건 버티기' 정도 (+로 클수록 습관 강함)",
       size=13, bold=True, color=NAVY, after=4)
    def fmt(x):
        return "측정 불가" if x is None else f"{x*100:+.0f}"
    _p(doc, f"  · 지지부진하던 장(2014–15) : {fmt(prof['de_box'])}",
       size=10.5, after=1)
    _p(doc, f"  · 쭉 오르던 장(2020–21)   : {fmt(prof['de_trend'])}",
       size=10.5, after=1)
    _p(doc, f"  · 전체                    : {fmt(prof['de_all'])}", size=10.5)

    # 손익 비교
    _p(doc, "■ 성과 비교 (같은 시장, 다른 습관)", size=13, bold=True, color=NAVY,
       after=4)
    for label, pct in extras["pnl_table"]:
        _p(doc, f"  · {label}: {pct:+.1f}%", size=10.5, after=1)
    if extras.get("missed_note"):
        _p(doc, f"  → {extras['missed_note']}", size=10.5, color=RED)

    # 정답 공개
    _p(doc, "■ 종목 정체 공개", size=13, bold=True, color=NAVY, after=4)
    for row in extras["reveal"]:
        _p(doc, f"  {row}", size=10, after=1)

    # 맞춤 처방
    _p(doc, "■ 나를 위한 맞춤 처방", size=13, bold=True, color=NAVY, after=4)
    for s in extras["solutions"]:
        _p(doc, f"  • {s}", size=10.5, after=3)

    # Shadow backtesting 요약(선택)
    if extras.get("shadow"):
        sh = extras["shadow"]
        inf = sh.get("inferred", {})
        summ = sh.get("summary", {})
        _p(doc, "■ 규칙대로 했다면? (내 매매 vs 규칙 매매)", size=13, bold=True,
           color=NAVY, after=4)
        _p(doc, f"  · 추정 매매 스타일: {inf.get('style', 'n/a')}",
           size=10.5, after=1)
        _p(doc, f"  · 규칙 지킨 비율: {summ.get('rule_adherence_pct', 0):.1f}%",
           size=10.5, after=1)
        _p(doc, f"  · 규칙 어긴 횟수: {summ.get('n_violations', 0)}건",
           size=10.5, after=1)
        _p(doc, f"  · 너무 일찍 판 횟수: "
           f"{summ.get('early_profit_take_count', 0)}건", size=10.5, after=1)
        _p(doc, f"  · 손절 늦은 횟수: "
           f"{summ.get('delayed_stop_loss_count', 0)}건", size=10.5, after=1)
        _p(doc, f"  · 손절 안 한 횟수: "
           f"{summ.get('no_stop_loss_execution_count', 0)}건",
           size=10.5, after=1)
        _p(doc, f"  · 놓친 수익 합계: "
           f"{summ.get('total_opportunity_cost_pct', 0):+.1f}%포인트", size=10.5)

    _p(doc, "본 진단은 단일 검사(10종목) 기반의 성향 지표이며, 정밀 측정이 "
            "아닙니다. 투자 권유가 아니며 판단과 책임은 본인에게 있습니다.",
       size=9, color=GRAY, after=2)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
